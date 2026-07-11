# --- 核心归档函数 ---
import os
import sqlite3
import shutil
from aurum_core.database import get_db_path
from aurum_core.text_utils import read_file_content, normalize_date_str, extract_diagnosis
import re

def clean_patient_name(name: str) -> str:
    """
    清洗患者姓名
    - 去除首尾空格
    - 去除括号及其内容（包括中文括号、英文括号）
    - 去除日期格式（如 2025-06-01）
    - 去除“初诊”、“复诊”、“首诊”等常见备注
    - 合并连续空格
    - 过滤非法文件名字符
    """
    if not name:
        return name

    # 1. 去除首尾空格
    name = name.strip()

    # 2. 去除括号及其内容（中文括号（）和英文括号()）
    name = re.sub(r'[（(][^）)]*[）)]', '', name)

    # 3. 去除“初诊”、“复诊”、“首诊”、“二诊”等
    name = re.sub(r'[初首二三四五六七八九]诊', '', name)

    # 4. 去除日期格式（如 2025-06-01、2025/06/01、2025年06月01日）
    name = re.sub(r'\d{4}[-/年]\d{1,2}[-/月]\d{1,2}日?', '', name)
    name = re.sub(r'\d{4}[-/]\d{1,2}[-/]\d{1,2}', '', name)

    # 5. 全角数字/字母转半角
    name = name.translate(str.maketrans(
        "０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ",
        "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz"
    ))

    # 6. 移除非法文件名字符（保留中文、英文、数字、下划线、连字符）
    name = re.sub(r'[<>:"/\\|?*]', '', name)

    # 7. 合并连续空格为一个
    name = re.sub(r'\s+', ' ', name)

    # 8. 去除首尾空格和多余标点
    name = name.strip(' .,;:！？。，；：')

    return name if name else "未知患者"

def clean_path(path: str) -> str:
    """去除路径字符串首尾的引号（双引号或单引号）"""
    if not path:
        return path
    return path.strip().strip('"').strip("'")

def select_medical_file(folder_path: str, expected_date: str = None) -> tuple:
    """
    智能选择患者文件夹中的病历文件（docx/doc/txt）。
    新增逻辑：
    1. 文件内容必须包含医学关键词（主诉、舌脉、诊断等）
    2. 文件内容必须包含就诊日期（如果提供了 expected_date）
    3. 支持多种日期格式解析（YYYY-MM-DD, YYYY/MM/DD, YYYY年MM月DD日等）
    返回 (文件路径, 状态信息) 其中状态信息为提示文本
    """
    import re
    from datetime import datetime

    # ---- 新增：检查文件夹是否存在 ----
    if not os.path.exists(folder_path):
        return None, f"文件夹不存在: {folder_path}"

    candidates = []
    for f in os.listdir(folder_path):
        file_path = os.path.join(folder_path, f)
        if not os.path.isfile(file_path):
            continue
        ext = os.path.splitext(f)[1].lower()
        if ext in ('.docx', '.doc', '.txt'):
            candidates.append(file_path)

    if not candidates:
        return None, "未找到任何 .docx/.doc/.txt 文件"

    # 解析期望日期
    target_date = None
    if expected_date:
        patterns = [
            r'(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})',
            r'(\d{4})年(\d{1,2})月(\d{1,2})日?',
            r'(\d{4})年(\d{1,2})月',
        ]
        for pat in patterns:
            m = re.search(pat, expected_date)
            if m:
                year = int(m.group(1))
                month = int(m.group(2))
                day = int(m.group(3)) if len(m.groups()) >= 3 else 1
                try:
                    target_date = datetime(year, month, day).date()
                    break
                except:
                    continue

    # 对每个候选文件，检查日期是否匹配
    date_matched_files = []
    for path in candidates:
        content = read_file_content(path)
        if not content:
            continue
        # 检查内容中是否包含目标日期（如果 target_date 存在）
        if target_date:
            found = False
            for pat in patterns:
                for m in re.finditer(pat, content):
                    try:
                        year = int(m.group(1))
                        month = int(m.group(2))
                        day = int(m.group(3)) if len(m.groups()) >= 3 else 1
                        found_date = datetime(year, month, day).date()
                        if found_date == target_date:
                            found = True
                            break
                    except:
                        continue
                if found:
                    break
            if found:
                date_matched_files.append(path)
        else:
            # 如果没有期望日期，则直接加入（但这种情况很少）
            date_matched_files.append(path)

    if date_matched_files:
        # 在日期匹配的文件中，优先选择包含医学关键词的
        for path in date_matched_files:
            content = read_file_content(path)
            if content and any(kw in content for kw in ['主诉', '舌脉', '诊断', '方药', '处方', '药方']):
                return path, "✅ 通过日期匹配且包含医学关键词"
        # 其次选文件最大的
        date_matched_files.sort(key=lambda x: os.path.getsize(x), reverse=True)
        return date_matched_files[0], "【病历文件】⚠️ 通过日期匹配但未找到明确医学关键词，请人工确认"
    else:
        # ---- 新增保底策略 ----
        # 如果没有任何文件匹配日期，但文件夹中有文档文件，尝试用保底逻辑
        if candidates:
            # 1. 优先选择包含医学关键词的文件（不检查日期）
            for path in candidates:
                content = read_file_content(path)
                if content and any(kw in content for kw in ['主诉', '舌脉', '诊断', '方药', '处方', '药方']):
                    return path, "【病历文件】⚠️ 未匹配日期，但包含医学关键词，已自动采用"
            # 2. 如果只有一个文档文件，直接采用
            if len(candidates) == 1:
                return candidates[0], "【病历文件】⚠️ 未匹配日期，且文件夹中仅有一个文档文件，已自动采用"
            # 3. 如果有多个文档，选文件最大的
            candidates.sort(key=lambda x: os.path.getsize(x), reverse=True)
            return candidates[0], "【病历文件】⚠️ 未匹配日期，已选文件最大的文档，请人工确认"
        else:
            return None, "❌ 未找到任何可识别的病历文件"

def _find_medical_file(folder_path: str) -> str:
    """在文件夹中查找第一个病历文件（.docx/.doc/.txt）"""
    for f in os.listdir(folder_path):
        if f.lower().endswith(('.docx', '.doc', '.txt')):
            return os.path.join(folder_path, f)
    return None

def _insert_visit_record(cursor, patient: str, visit_date: str, hospital: str, docx_path: str):
    if not os.path.exists(docx_path):
        return False
    if os.path.getsize(docx_path) == 0:
        return False

    try:
        content = read_file_content(docx_path)
    except (PermissionError, OSError):
        # 文件被占用（如 Word 打开），返回特殊标记
        return "occupied"
    except Exception:
        return False

    if not content or not content.strip():
        return False

    # 尝试读取文件内容，捕获权限错误
    try:
        content = read_file_content(docx_path)
    except PermissionError:
        # 文件被占用（如 Word 打开），返回 False 并记录特殊日志
        return "occupied"
    except Exception:
        return False

    if not content or not content.strip():
        return False

    try:
        cursor.execute(
            "SELECT id FROM visits WHERE patient_name = ? AND visit_date = ? AND hospital = ?",
            (patient, visit_date, hospital)
        )
        existing = cursor.fetchone()

        if existing:
            cursor.execute(
                "UPDATE visits SET docx_path = ?, full_medical_text = ? WHERE patient_name = ? AND visit_date = ? AND hospital = ?",
                (docx_path, content, patient, visit_date, hospital)
            )
        else:
            cursor.execute('''
                INSERT INTO visits 
                (patient_name, visit_date, hospital, docx_path, full_medical_text)
                VALUES (?, ?, ?, ?, ?)
            ''', (patient, visit_date, hospital, docx_path, content))

        from aurum_core.text_utils import extract_diagnosis
        tcm_list, wm_list = extract_diagnosis(content, default_to_tcm=True)
        import json
        if tcm_list:
            cursor.execute(
                "UPDATE visits SET diagnosis = ? WHERE patient_name = ? AND visit_date = ? AND hospital = ?",
                (json.dumps(tcm_list, ensure_ascii=False), patient, visit_date, hospital)
            )
        if wm_list:
            cursor.execute(
                "UPDATE visits SET western_diagnosis = ? WHERE patient_name = ? AND visit_date = ? AND hospital = ?",
                (json.dumps(wm_list, ensure_ascii=False), patient, visit_date, hospital)
            )
        return True
    except Exception:
        return False

def reorganize_files(src, tgt, aggregate: bool = False):
    # --- 第一步：检查源目录 ---
    if not os.path.exists(src):
        return "❌ 源路径不存在，请检查"

    try:
        items = os.listdir(src)
        if not items:
            return "⚠️ 源目录为空，未找到任何文件夹，请检查路径是否正确。"
    except PermissionError:
        return "⛔ 没有权限读取源文件夹"

    # ---- 日期格式匹配 ----
    date_pattern = re.compile(
        r'^(\d{4}[-/.]\d{1,2}[-/.]\d{1,2}|\d{4}年\d{1,2}月\d{1,2}日?)(?:\s+(.*))?$'
    )

    # ---- 先检测根目录类型（不创建目标目录） ----
    base_name = os.path.basename(src)
    match_root = date_pattern.match(base_name)
    log_lines = []

    if match_root and os.path.isdir(src):
        # 根目录本身就是日期文件夹 → 结构A
        date_folders = [base_name]
        root_is_date = True
        log_lines.append(f"🔍 根目录：{src}，发现根目录为日期文件夹")
    else:
        # 检查是否有日期子文件夹
        date_folders = []
        for f in items:
            full_path = os.path.join(src, f)
            if os.path.isdir(full_path) and date_pattern.match(f):
                date_folders.append(f)
        if date_folders:
            root_is_date = False
            log_lines.append(f"🔍 根目录：{src}，发现 {len(date_folders)} 个日期子文件夹")
        else:
            # 没有日期文件夹，检查是否为已归档结构（医院/患者/日期）→ 结构B
            hospital_candidates = [f for f in items if os.path.isdir(os.path.join(src, f))]
            if hospital_candidates:
                return (
                    "❌ 检测到您输入的是已整理好的归档文件夹（医院→患者→日期），请使用的「🔄 刷新数据库」折叠面板"
                )
            else:
                return "⚠️ 归档失败：根目录下既没有日期文件夹，也没有可识别的医院文件夹。请确认源目录是否正确。"

    # ========== 以下为结构A（日期文件夹，需要复制） ==========
    # ---- 检查源目录与目标目录是否相同（仅结构A需要） ----
    if os.path.abspath(src) == os.path.abspath(tgt):
        return "⚠️ 源目录与目标目录相同，请修改目标目录为不同路径，否则可能导致文件混乱。"

    # ---- 创建目标目录（结构A专用） ----
    try:
        os.makedirs(tgt, exist_ok=True)
    except Exception as e:
        return f"❌ 目标目录路径无效或无法创建：{e}"

    # ---- 结构A处理流程 ----
    db_path = get_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    patient_data = {}
    total_patients = 0

    # 统计总记录数
    total_visits = 0
    for date_folder in date_folders:
        if root_is_date:
            date_path = src
        else:
            date_path = os.path.join(src, date_folder)
        sub_items = os.listdir(date_path)
        patient_folders = [p for p in sub_items if os.path.isdir(os.path.join(date_path, p))]
        total_visits += len(patient_folders)
    log_lines.append(f"📋 共需处理 {total_visits} 条就诊记录")

    # ---- 第一阶段：复制并写入数据库 ----
    for date_folder in date_folders:
        if root_is_date:
            date_path = src
        else:
            date_path = os.path.join(src, date_folder)

        m = date_pattern.match(date_folder)
        if not m:
            log_lines.append(f"   ⚠️ 跳过无法解析的日期文件夹：{date_folder}")
            continue
        raw_date = m.group(1)
        hospital_name = m.group(2) if m.group(2) else "未归类患者"
        pure_date = normalize_date_str(raw_date)

        sub_items = os.listdir(date_path)
        patient_folders = [p for p in sub_items if os.path.isdir(os.path.join(date_path, p))]

        if not patient_folders:
            log_lines.append(f"   ⚠️ {date_folder} 下没有患者文件夹，跳过")
            continue

        for patient_raw in patient_folders:
            patient = clean_patient_name(patient_raw)
            patient_src_path = os.path.join(date_path, patient_raw)
            patient_tgt_path = os.path.join(tgt, hospital_name, patient, pure_date)

            if patient not in patient_data:
                patient_data[patient] = {
                    'gender': None,
                    'birth_date': None,
                    'phone': None,
                    'id_card': None,
                    'address': None,
                    'visits': []
                }
            patient_data[patient]['visits'].append({
                'date': pure_date,
                'hospital': hospital_name,
                'diagnosis': '',
                'prescription': '',
                'remarks': ''
            })

            file_count = 0
            for root, dirs, files in os.walk(patient_src_path):
                file_count += len(files)

            if file_count == 0:
                log_lines.append(f"   ⚠️ {patient}（{pure_date}，{hospital_name}）文件夹为空，跳过")
                continue

            shutil.copytree(patient_src_path, patient_tgt_path, dirs_exist_ok=True)
            total_patients += 1

            # 使用 select_medical_file 智能选择病历文件
            docx_file, select_status = select_medical_file(patient_tgt_path, pure_date)
            if select_status:
                log_lines.append(f"   ℹ️ {patient}（{pure_date}）{select_status}")

            if docx_file:
                result = _insert_visit_record(cursor, patient, pure_date, hospital_name, docx_file)
                if result is True:
                    total_patients += 1
                elif result == "occupied":
                    log_lines.append(f"   ⚠️ {patient}（{pure_date}）病历文件被占用，请关闭后重试")
                else:
                    log_lines.append(f"   ⚠️ {patient}（{pure_date}）病历文件无效或为空，已跳过")
            else:
                log_lines.append(f"   ⚠️ {patient}（{pure_date}，{hospital_name}）未找到病历文件（.docx/.doc/.txt），已跳过")

    conn.commit()

    # ---- 提取个人信息并生成档案（结构A专用） ----
    if total_patients > 0:
        try:
            from aurum_core.extract_patient_profiles import update_all_profiles
            update_all_profiles()
            log_lines.append("   ✅ 患者个人信息已提取（性别/出生日期/电话等）")
        except Exception as e:
            log_lines.append(f"   ⚠️ 自动提取个人信息失败：{e}")

        from aurum_core.file_processor import update_patient_archive_by_db
        for patient in patient_data.keys():
            try:
                update_patient_archive_by_db(patient, aggregate=aggregate)
                log_lines.append(f"   ✅ {patient} 的档案已生成")
            except Exception as e:
                log_lines.append(f"   ⚠️ {patient} 档案生成失败：{e}")

    conn.close()
    log_lines.append(f"\n🎉 归档完成！请前往 `{tgt}` 查看整理后的文件。")
    return "\n".join(log_lines)

def reorganize_single_hospital(hospital_path: str, tgt_root: str, hospital_name: str) -> str:
    """
    处理单个医院文件夹，支持两种结构：
    1. 医院/日期/患者（未归档） → 复制文件到目标目录，并更新数据库 + 生成档案
    2. 医院/患者/日期（已归档） → 只更新数据库，完全不碰目标目录（不创建任何文件夹或文件）
    """
    import shutil
    from aurum_core.text_utils import read_file_content, normalize_date_str, extract_diagnosis

    if not os.path.exists(hospital_path):
        return f"❌ 医院路径不存在：{hospital_path}"

    date_pattern = re.compile(
        r'^(\d{4}[-/.]\d{1,2}[-/.]\d{1,2}|\d{4}年\d{1,2}月\d{1,2}日?)(?:\s+(.*))?$'
    )

    db_path = get_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    log_lines = []
    total_success = 0
    total_failed = 0
    total_skipped = 0
    updated_patients = []   # 仅用于结构A

    items = os.listdir(hospital_path)
    subdirs = [f for f in items if os.path.isdir(os.path.join(hospital_path, f))]

    if not subdirs:
        return f"⚠️ 医院 [{hospital_name}] 下没有子目录"

    date_folders = [f for f in subdirs if date_pattern.match(f)]

    # ========== 结构A：医院/日期/患者（复制模式） ==========
    if date_folders:
        log_lines.append(f"🏥 处理科室：{hospital_name}（日期→患者结构，【复制模式】）")
        for date_folder in date_folders:
            date_path = os.path.join(hospital_path, date_folder)
            m = date_pattern.match(date_folder)
            if not m:
                continue
            raw_date = m.group(1)
            pure_date = normalize_date_str(raw_date)

            patient_folders = [p for p in os.listdir(date_path) if os.path.isdir(os.path.join(date_path, p))]
            for patient_raw in patient_folders:
                patient = clean_patient_name(patient_raw)
                patient_src_path = os.path.join(date_path, patient_raw)
                patient_tgt_path = os.path.join(tgt_root, hospital_name, patient, pure_date)

                # 检查是否为空文件夹
                file_count = sum(len(files) for _, _, files in os.walk(patient_src_path))
                if file_count == 0:
                    log_lines.append(f"   ⚠️ {patient}（{pure_date}）文件夹为空，跳过")
                    total_skipped += 1
                    continue

                # 复制整个文件夹
                try:
                    shutil.copytree(patient_src_path, patient_tgt_path, dirs_exist_ok=True)
                except Exception as e:
                    log_lines.append(f"   ❌ {patient}（{pure_date}）复制失败：{e}")
                    total_failed += 1
                    continue

                # 查找病历文件
                docx_file, select_status = select_medical_file(patient_tgt_path, pure_date)
                if select_status:
                    log_lines.append(f"   ℹ️ {patient}（{pure_date}）{select_status}")

                if docx_file:
                    result = _insert_visit_record(cursor, patient, pure_date, hospital_name, docx_file)
                    if result is True:
                        total_success += 1
                        updated_patients.append(patient)
                    elif result == "occupied":
                        log_lines.append(f"   ⚠️ {patient}（{pure_date}）病历文件被占用，请关闭后重试")
                        total_failed += 1
                    else:
                        log_lines.append(f"   ⚠️ {patient}（{pure_date}）病历文件无效或为空，已跳过")
                        total_failed += 1
                else:
                    log_lines.append(f"   ⚠️ {patient}（{pure_date}）未找到病历文件，已跳过")
                    total_failed += 1

        conn.commit()
        conn.close()

        # ---- 结构A专有：提取个人信息 + 生成档案 ----
        if updated_patients:
            try:
                from aurum_core.extract_patient_profiles import update_all_profiles
                update_all_profiles()
                log_lines.append("   ✅ 患者个人信息已提取（性别/出生日期/电话等）")
            except Exception as e:
                log_lines.append(f"   ⚠️ 自动提取个人信息失败：{e}")

            from aurum_core.file_processor import update_patient_archive_by_db
            for patient in set(updated_patients):
                try:
                    update_patient_archive_by_db(patient, aggregate=False)
                    log_lines.append(f"   ✅ {patient} 的档案已生成/更新")
                except Exception as e:
                    log_lines.append(f"   ⚠️ 生成档案失败（{patient}）：{e}")

        log_lines.append(f"📊 {hospital_name}：成功 {total_success} 位，失败 {total_failed} 位，跳过 {total_skipped} 个空文件夹")
        return "\n".join(log_lines)

    # ========== 结构B：医院/患者/日期（【仅刷新数据库，绝不创建任何文件夹】） ==========
    else:
        log_lines.append(f"🏥 处理科室：{hospital_name}（患者→日期结构，【仅刷新数据库，不创建任何文件夹】）")
        for patient_raw in subdirs:
            patient = clean_patient_name(patient_raw)
            patient_path = os.path.join(hospital_path, patient_raw)
            # 获取该患者下的所有日期子文件夹
            date_subdirs = [f for f in os.listdir(patient_path)
                            if os.path.isdir(os.path.join(patient_path, f)) and date_pattern.match(f)]
            if not date_subdirs:
                log_lines.append(f"   ⚠️ {patient} 文件夹下没有日期子文件夹，跳过")
                total_skipped += 1
                continue

            for date_folder in date_subdirs:
                m = date_pattern.match(date_folder)
                if not m:
                    continue
                raw_date = m.group(1)
                pure_date = normalize_date_str(raw_date)
                patient_date_path = os.path.join(patient_path, date_folder)

                # 检查是否为空
                file_count = sum(len(files) for _, _, files in os.walk(patient_date_path))
                if file_count == 0:
                    log_lines.append(f"   ⚠️ {patient}（{pure_date}）文件夹为空，跳过")
                    total_skipped += 1
                    continue

                # 查找病历文件（源路径）
                docx_file, select_status = select_medical_file(patient_date_path, pure_date)
                if select_status:
                    log_lines.append(f"   ℹ️ {patient}（{pure_date}）{select_status}")

                if docx_file:
                    result = _insert_visit_record(cursor, patient, pure_date, hospital_name, docx_file)
                    if result is True:
                        total_success += 1
                    elif result == "occupied":
                        log_lines.append(f"   ⚠️ {patient}（{pure_date}）病历文件被占用，请关闭后重试")
                        total_failed += 1
                    else:
                        log_lines.append(f"   ⚠️ {patient}（{pure_date}）病历文件无效或为空，已跳过")
                        total_failed += 1
                else:
                    log_lines.append(f"   ⚠️ {patient}（{pure_date}）未找到病历文件，已跳过")
                    total_failed += 1

        conn.commit()
        conn.close()
        log_lines.append(f"📊 {hospital_name}：成功更新 {total_success} 条记录，失败 {total_failed} 条，跳过 {total_skipped} 个空文件夹")
        log_lines.append("✅ 处理完成：未在目标目录创建任何文件或文件夹，仅数据库已同步。")
        return "\n".join(log_lines)

def read_word_content(doc_path):
    from docx import Document
    doc = Document(doc_path)
    text = []
    for para in doc.paragraphs:
        if para.text:
            text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text.append(cell.text)
    return "\n".join(text)

def reindex_archived_folder(tgt_root: str) -> str:
    """
    重建索引：扫描归档文件夹，更新 visits 表中的 docx_path。
    不覆盖诊断、方剂等已有数据，只更新文件路径。
    """
    if not os.path.exists(tgt_root):
        return "❌ 路径不存在，请检查"

    try:
        sub_items = os.listdir(tgt_root)
        has_hospital_folder = any(os.path.isdir(os.path.join(tgt_root, f)) for f in sub_items)
        if not has_hospital_folder:
            return "⚠️ 路径下没有找到医院/科室文件夹，请确保路径指向归档根目录（如 D:/test_Aurum归档）"
    except PermissionError:
        return "⛔ 没有权限读取该文件夹"

    db_path = get_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    total_updated = 0
    total_missing = 0
    log_lines = []
    log_lines.append(f"🔍 开始重建索引，根目录：{tgt_root}")

    for hospital in os.listdir(tgt_root):
        hospital_path = os.path.join(tgt_root, hospital)
        if not os.path.isdir(hospital_path):
            continue

        for patient in os.listdir(hospital_path):
            patient_path = os.path.join(hospital_path, patient)
            if not os.path.isdir(patient_path):
                continue

            for date_folder in os.listdir(patient_path):
                date_path = os.path.join(patient_path, date_folder)
                if not os.path.isdir(date_path):
                    continue

                docx_file = None
                for f in os.listdir(date_path):
                    if f.lower().endswith(('.docx', '.doc', '.txt')):
                        docx_file = os.path.join(date_path, f)
                        break

                if not docx_file:
                    log_lines.append(f"   ⚠️ 缺失病历：{patient}/{date_folder}")
                    total_missing += 1
                    continue

                cursor.execute(
                    "UPDATE visits SET docx_path = ? WHERE patient_name = ? AND visit_date = ?",
                    (docx_file, patient, date_folder)
                )
                if cursor.rowcount > 0:
                    total_updated += 1
                else:
                    cursor.execute('''
                        INSERT OR IGNORE INTO visits (patient_name, visit_date, hospital, docx_path)
                        VALUES (?, ?, ?, ?)
                    ''', (patient, date_folder, hospital, docx_file))
                    if cursor.rowcount > 0:
                        total_updated += 1

    conn.commit()
    log_lines.append(f"\n📊 重建索引完成！共更新 {total_updated} 条记录，缺失 {total_missing} 个病历文件。")
    return "\n".join(log_lines)

def update_patient_archive_by_db(patient_name: str, aggregate: bool = False):
    """
    根据数据库记录，自动发现所有归档根目录并更新 _患者档案.txt。
    用于编辑个人信息后同步更新所有根目录下的档案。

    参数:
        patient_name: 患者姓名
        aggregate: 是否汇总所有根目录的就诊记录
            - False（默认）：每个根目录下的档案只包含该目录内的就诊记录
            - True：每个根目录下的档案都包含该患者的全部就诊记录
    """
    import sqlite3
    import json
    from datetime import datetime
    from collections import defaultdict

    db_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "aurum_index.db")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 获取该患者所有就诊记录（含 docx_path）
    cursor.execute(
        "SELECT visit_date, hospital, diagnosis, prescription, visit_remarks, docx_path "
        "FROM visits WHERE patient_name = ? ORDER BY visit_date ASC",
        (patient_name,)
    )
    rows = cursor.fetchall()

    if not rows:
        conn.close()
        print(f"⚠️ 患者 {patient_name} 没有就诊记录，无法生成档案")
        return

    # 获取患者档案信息
    cursor.execute(
        "SELECT gender, birth_date, phone, id_card, address, personal_remarks "
        "FROM patient_profiles WHERE patient_name = ?",
        (patient_name,)
    )
    profile = cursor.fetchone()
    conn.close()

    # 按归档根目录分组，同时收集全部记录（用于汇总模式）
    root_to_rows = defaultdict(list)
    all_rows = []
    for row in rows:
        visit_date, hospital, diagnosis, prescription, remarks, docx_path = row
        if docx_path and os.path.exists(docx_path):
            # 提取根目录：.../根目录/医院/患者/日期/病历.docx
            date_folder = os.path.dirname(docx_path)
            patient_folder = os.path.dirname(date_folder)  # .../根目录/医院/患者
            hospital_folder = os.path.dirname(patient_folder)  # .../根目录/医院
            root = os.path.dirname(hospital_folder)  # .../根目录
            root_to_rows[root].append(row)
            all_rows.append(row)
        else:
            # 如果 docx_path 无效，跳过（这种情况极少）
            continue

    if not root_to_rows:
        print(f"⚠️ 患者 {patient_name} 的所有记录都没有有效的文件路径")
        return

    # ---- 根据 aggregate 决定使用哪组数据 ----
    # 如果开启汇总，所有根目录共用 all_rows；否则各自使用自己的 rows_in_root
    for root, rows_in_root in root_to_rows.items():
        # 选择数据源
        data_source = all_rows if aggregate else rows_in_root

        lines = []
        lines.append("患者档案")
        lines.append("")
        lines.append(f"姓名：{patient_name}")
        lines.append(f"性别：{profile[0] if profile else '无'}")
        lines.append(f"出生日期：{profile[1] if profile else '无'}")
        lines.append(f"电话：{profile[2] if profile else '无'}")
        lines.append(f"身份证号：{profile[3] if profile else '无'}")
        lines.append(f"住址：{profile[4] if profile else '无'}")
        if profile and profile[5]:
            lines.append(f"个人信息备注：{profile[5]}")
        lines.append("")
        lines.append("【就诊记录汇总】")

        for row in data_source:
            visit_date, hospital, diagnosis, prescription, remarks, _ = row
            try:
                diag_list = json.loads(diagnosis) if diagnosis else []
                diag_str = '、'.join(diag_list) if diag_list else '待补充'
            except:
                diag_str = diagnosis or '待补充'
            try:
                presc_list = json.loads(prescription) if prescription else []
                presc_str = '、'.join(presc_list) if presc_list else ''
            except:
                presc_str = prescription or ''
            line = f" {hospital} | {visit_date} | 诊断：{diag_str}"
            if presc_str:
                line += f" | 方剂：{presc_str}"
            if remarks:
                line += f" | 备注：{remarks}"
            lines.append(line)

        lines.append("")
        if aggregate:
            lines.append(
                f"共 {len(all_rows)} 次就诊记录（汇总自所有归档目录）| 档案更新于：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        else:
            lines.append(f"共 {len(rows_in_root)} 次就诊记录 | 档案更新于：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

        # 找到该患者在该根目录下的所有医院文件夹（通过遍历 rows_in_root）
        for row in rows_in_root:
            _, hospital, _, _, _, docx_path = row
            if docx_path and os.path.exists(docx_path):
                patient_dir = os.path.join(root, hospital, patient_name)
                archive_path = os.path.join(patient_dir, "_患者档案.txt")
                os.makedirs(patient_dir, exist_ok=True)
                try:
                    with open(archive_path, 'w', encoding='utf-8') as f:
                        f.write('\n'.join(lines))
                    print(f"✅ 已更新档案：{archive_path}")
                except Exception as e:
                    print(f"❌ 写入档案失败：{archive_path}, {e}")

def refresh_index_only(src_root: str, aggregate: bool = False) -> str:
    """
    仅刷新数据库索引，不复制任何文件，不创建任何文件夹。
    用于已有的归档结构（医院/患者/日期），将最新文件路径和内容同步到数据库。
    完成后会同步更新所有患者的档案文件。
    """
    if not os.path.exists(src_root):
        return "❌ 源路径不存在，请检查"

    try:
        items = os.listdir(src_root)
        if not items:
            return "⚠️ 源目录为空"
    except PermissionError:
        return "⛔ 没有权限读取源文件夹"

    # ---- 检测是否为结构A（日期文件夹） ----
    date_pattern = re.compile(
        r'^(\d{4}[-/.]\d{1,2}[-/.]\d{1,2}|\d{4}年\d{1,2}月\d{1,2}日?)(?:\s+(.*))?$'
    )
    base_name = os.path.basename(src_root)
    if date_pattern.match(base_name):
        return (
            "❌ 检测到日期文件夹（未整理病历），请使用主界面的「归档整理」功能进行复制归档。\n"
            "   刷新数据库仅适用于已整理好的归档文件夹（医院/患者/日期）。"
        )
    date_subdirs = [f for f in items if os.path.isdir(os.path.join(src_root, f)) and date_pattern.match(f)]
    hospital_candidates = [f for f in items if os.path.isdir(os.path.join(src_root, f)) and not date_pattern.match(f)]
    if date_subdirs and not hospital_candidates:
        return (
            "❌ 检测到日期文件夹集合（未整理病历），请使用主界面的「归档整理」功能进行复制归档。\n"
            "   刷新数据库仅适用于已整理好的归档文件夹（医院/患者/日期）。"
        )

    log_lines = []
    log_lines.append(f"🔄 开始刷新数据库索引（仅更新，不复制文件）...")
    log_lines.append(f"📂 根目录：{src_root}")

    db_path = get_db_path()
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    total_updated = 0
    total_inserted = 0
    total_errors = 0
    total_skipped_no_file = 0
    total_skipped_empty = 0
    total_skipped_occupied = 0

    # ---- 直接扫描三层结构：医院/患者/日期 ----
    for hospital in os.listdir(src_root):
        hospital_path = os.path.join(src_root, hospital)
        if not os.path.isdir(hospital_path):
            continue
        log_lines.append(f"🏥 处理医院：{hospital}")
        for patient in os.listdir(hospital_path):
            patient_path = os.path.join(hospital_path, patient)
            if not os.path.isdir(patient_path):
                continue
            log_lines.append(f"   👤 患者：{patient}")
            for date_folder in os.listdir(patient_path):
                date_path = os.path.join(patient_path, date_folder)
                if not os.path.isdir(date_path):
                    continue
                pure_date = normalize_date_str(date_folder)
                if not pure_date:
                    log_lines.append(f"      ⚠️ 日期格式无法解析：{date_folder}，跳过")
                    total_errors += 1
                    continue

                # 查找第一个文档文件
                docx_file = None
                for f in os.listdir(date_path):
                    f_path = os.path.join(date_path, f)
                    if os.path.isfile(f_path) and f.lower().endswith(('.docx', '.doc', '.txt')):
                        docx_file = f_path
                        break
                if not docx_file:
                    log_lines.append(f"      ⚠️ {date_folder} 未找到文档文件，跳过")
                    total_skipped_no_file += 1
                    continue

                # 读取文件内容，区分具体错误
                try:
                    content = read_file_content(docx_file)
                except (PermissionError, OSError):
                    log_lines.append(f"      ⚠️ 文件被占用（可能正在编辑）：{docx_file}，跳过")
                    total_skipped_occupied += 1
                    total_errors += 1
                    continue
                except Exception as e:
                    log_lines.append(f"      ❌ 读取文件失败：{docx_file}，错误：{e}")
                    total_errors += 1
                    continue

                if not content or not content.strip():
                    log_lines.append(f"      ⚠️ 文件内容为空：{docx_file}，跳过")
                    total_skipped_empty += 1
                    total_errors += 1
                    continue

                # 检查记录是否存在
                cursor.execute(
                    "SELECT id FROM visits WHERE patient_name = ? AND visit_date = ? AND hospital = ?",
                    (patient, pure_date, hospital)
                )
                existing = cursor.fetchone()
                if existing:
                    cursor.execute(
                        "UPDATE visits SET docx_path = ?, full_medical_text = ? WHERE id = ?",
                        (docx_file, content, existing[0])
                    )
                    total_updated += 1
                    log_lines.append(f"      ✅ 更新记录：{patient} | {pure_date} | {hospital}")
                else:
                    cursor.execute(
                        "INSERT INTO visits (patient_name, visit_date, hospital, docx_path, full_medical_text) VALUES (?, ?, ?, ?, ?)",
                        (patient, pure_date, hospital, docx_file, content)
                    )
                    total_inserted += 1
                    log_lines.append(f"      ✅ 新增记录：{patient} | {pure_date} | {hospital}")

                # 提取诊断
                tcm_list, wm_list = extract_diagnosis(content, default_to_tcm=True)
                if tcm_list:
                    cursor.execute(
                        "UPDATE visits SET diagnosis = ? WHERE patient_name = ? AND visit_date = ? AND hospital = ?",
                        (json.dumps(tcm_list, ensure_ascii=False), patient, pure_date, hospital)
                    )
                if wm_list:
                    cursor.execute(
                        "UPDATE visits SET western_diagnosis = ? WHERE patient_name = ? AND visit_date = ? AND hospital = ?",
                        (json.dumps(wm_list, ensure_ascii=False), patient, pure_date, hospital)
                    )

    conn.commit()
    conn.close()

    # ---- 统计摘要 ----
    log_lines.append("")
    log_lines.append(f"📊 刷新统计：")
    log_lines.append(f"   - 更新记录数：{total_updated}")
    log_lines.append(f"   - 新增记录数：{total_inserted}")
    log_lines.append(f"   - 跳过（无文档）：{total_skipped_no_file}")
    log_lines.append(f"   - 跳过（内容为空）：{total_skipped_empty}")
    log_lines.append(f"   - 跳过（文件被占用）：{total_skipped_occupied}")
    log_lines.append(f"   - 总错误数：{total_errors}")

    # ---- 提取个人信息 ----
    try:
        from aurum_core.extract_patient_profiles import update_all_profiles
        update_all_profiles()
        log_lines.append("✅ 患者个人信息已提取（性别/出生日期/电话等）")
    except Exception as e:
        log_lines.append(f"⚠️ 自动提取个人信息失败：{e}")

    # ---- 同步更新所有患者的档案文件 ----
    try:
        from aurum_core.file_processor import update_patient_archive_by_db
        conn2 = sqlite3.connect(db_path)
        cursor2 = conn2.cursor()
        cursor2.execute("SELECT DISTINCT patient_name FROM visits")
        patients = [row[0] for row in cursor2.fetchall()]
        conn2.close()
        if patients:
            for patient in patients:
                try:
                    update_patient_archive_by_db(patient, aggregate=aggregate)
                except Exception as e:
                    log_lines.append(f"⚠️ 生成档案失败（{patient}）：{e}")
            log_lines.append(f"✅ 已同步更新 {len(patients)} 位患者的档案文件")
    except Exception as e:
        log_lines.append(f"⚠️ 更新档案文件失败：{e}")

    log_lines.append(f"✅ 刷新完成：更新 {total_updated} 条，新增 {total_inserted} 条，错误 {total_errors} 条。")
    return "\n".join(log_lines)